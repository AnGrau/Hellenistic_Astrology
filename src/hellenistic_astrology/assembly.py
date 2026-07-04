"""Assemble le document final : ajoute la prose de Phase 3 (rédigée en
session assistée, jalons 26-27, sous forme de texte Markdown) à la suite
d'un `.docx` de Phase 1/2 déjà généré (`docgen.builder`).

Nature de tâche distincte de `docgen` (qui part toujours d'un `Observation`
frais, jamais d'un `.docx` déjà existant sur disque) et d'`interpretation`
(qui ne produit qu'un brief, jamais de document). Aucun calcul
astrologique ici : le texte Markdown est traité tel quel, sans vérifier
son contenu — cette vérification a déjà eu lieu pendant la rédaction
assistée elle-même.

Le sous-ensemble Markdown géré est volontairement étroit, borné à ce que
produit effectivement le skill de rédaction (jalon 27) : titres `#`/`##`,
paragraphes de texte brut, puces `- `, et emphase `*texte*` (italique).
Pas de gras, pas de code inline, pas de listes numérotées — pas de
dépendance Markdown générale nécessaire pour ce périmètre précis.
"""

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document

_EMPHASIS_RE = re.compile(r"\*([^*]+)\*")


@dataclass(frozen=True)
class MarkdownBlock:
    kind: str  # "heading1" | "heading2" | "paragraph" | "bullet"
    segments: list[tuple[str, bool]]  # (texte, italique)


def _parse_segments(text: str) -> list[tuple[str, bool]]:
    """Découpe `text` en segments (texte, italique) sur la seule emphase
    gérée (`*...*`) — pas de gras, pas d'imbrication."""
    segments = []
    pos = 0
    for match in _EMPHASIS_RE.finditer(text):
        if match.start() > pos:
            segments.append((text[pos : match.start()], False))
        segments.append((match.group(1), True))
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], False))
    return segments or [("", False)]


def parse_markdown_blocks(text: str) -> list[MarkdownBlock]:
    """Découpe un texte Markdown (sous-ensemble étroit, voir docstring de
    module) en blocs typés, une ligne non vide = un bloc : les lignes
    vides ne sont que des séparateurs, jamais des blocs en elles-mêmes
    (cohérent avec la production du skill de rédaction, jalon 27, qui
    n'émet jamais de paragraphe étalé sur plusieurs lignes)."""
    blocks = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            kind, content = "heading2", line[3:]
        elif line.startswith("# "):
            kind, content = "heading1", line[2:]
        elif line.startswith("- "):
            kind, content = "bullet", line[2:]
        else:
            kind, content = "paragraph", line
        blocks.append(MarkdownBlock(kind=kind, segments=_parse_segments(content)))
    return blocks


def _plain_text(segments: list[tuple[str, bool]]) -> str:
    return "".join(text for text, _italic in segments)


def _add_paragraph_with_segments(document: Document, segments: list[tuple[str, bool]], style: str | None) -> None:
    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    for text, italic in segments:
        run = paragraph.add_run(text)
        if italic:
            run.italic = True


def append_markdown_to_document(document: Document, markdown_text: str) -> None:
    """Rejoue les blocs de `markdown_text` sur un `Document` déjà chargé —
    titres en `add_heading`, puces et paragraphes en runs mixtes
    (normal/italique) pour préserver l'emphase."""
    for block in parse_markdown_blocks(markdown_text):
        if block.kind == "heading1":
            document.add_heading(_plain_text(block.segments), level=1)
        elif block.kind == "heading2":
            document.add_heading(_plain_text(block.segments), level=2)
        elif block.kind == "bullet":
            _add_paragraph_with_segments(document, block.segments, style="List Bullet")
        else:
            _add_paragraph_with_segments(document, block.segments, style=None)


def assemble_final_document(docx_path: str | Path, markdown_path: str | Path, output_path: str | Path | None = None) -> Path:
    """Charge `docx_path` (déjà généré, Phase 1/2), y ajoute le contenu de
    `markdown_path` (Phase 3 déjà rédigée et finalisée), et écrit vers
    `output_path` — par défaut `<docx_path>_final.docx`, à côté de
    l'original. Ne modifie jamais `docx_path` sur place (cohérent avec la
    consigne de ne jamais écraser silencieusement un fichier existant)."""
    docx_path = Path(docx_path)
    markdown_path = Path(markdown_path)
    if not docx_path.exists():
        raise ValueError(f"fichier .docx introuvable : {docx_path}")
    if not markdown_path.exists():
        raise ValueError(f"fichier Markdown introuvable : {markdown_path}")

    document = Document(str(docx_path))
    append_markdown_to_document(document, markdown_path.read_text(encoding="utf-8"))

    target = Path(output_path) if output_path else docx_path.with_name(f"{docx_path.stem}_final.docx")
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))
    return target
