from datetime import datetime

from docx import Document

from ..core.aspects import ClusterAspect, SignCluster, sign_aspect
from ..core.dignities import DOMICILES, MutualReception
from ..core.houses import house_quality, index_of_sign
from ..core.observation import Observation, PointPosition
from ..core.zodiacal_releasing import ReleasingChapter, ReleasingPeriod, is_peak_period
from . import styles

FEMININE_PLANETS = {"Lune", "Vénus"}
FEMININE_POINTS = FEMININE_PLANETS | {"Part de Fortune", "Part de l'Esprit", "Part d'Éros"}

# Noms affichés avec article, pour les puces d'aspects (les planètes et les
# nœuds s'affichent sans article, comme dans les documents de référence).
DISPLAY_NAME_WITH_ARTICLE = {
    "Ascendant": "l'Ascendant",
    "Milieu du Ciel": "le Milieu du Ciel",
    "Part de Fortune": "la Part de Fortune",
    "Part de l'Esprit": "la Part de l'Esprit",
    "Part d'Éros": "la Part d'Éros",
}

ASPECT_LABEL = {
    "Sextile": "sextile",
    "Carré": "carré",
    "Trigone": "trigone",
    "Opposition": "opposition",
    "Aversion": "aversion",
}

POSITIONS_HEADER = ["Astre", "Signe", "Degré", "Maison", "Rôle de secte", "Direction", "Dignité essentielle"]
POSITIONS_COLUMN_WIDTHS_DXA = [1700, 1300, 1100, 900, 1900, 1300, 1600]

RULERSHIPS_HEADER = ["Planète", "Domiciles gouvernés", "Maisons régies depuis l'Ascendant"]
RULERSHIPS_COLUMN_WIDTHS_DXA = [2400, 3200, 3200]

MINOR_DIGNITIES_HEADER = ["Astre", "Triplicité", "Terme (bornes égyptiennes)", "Décan"]
MINOR_DIGNITIES_COLUMN_WIDTHS_DXA = [1700, 2900, 2900, 2300]

ZODIACAL_RELEASING_HEADER = ["Niveau", "Signe", "Maître", "Début", "Fin", "Culminante"]
ZODIACAL_RELEASING_COLUMN_WIDTHS_DXA = [900, 1700, 1700, 1600, 1600, 1400]

# Ordre confirmé par les deux documents de référence (Pérégrin traité à part,
# voir add_dignities_and_receptions_section) ; Domicile jamais illustré dans
# les deux thèmes, ordre déduit par cohérence avec le reste de la séquence.
DIGNITY_CATEGORY_ORDER = ["Domicile", "Chute", "Exaltation", "Exil (détriment)"]
DIGNITY_LABELS = {
    "Domicile": "En domicile",
    "Chute": "En chute",
    "Exaltation": "En exaltation",
    "Exil (détriment)": "En exil (détriment)",
}
# Seuil « sous les rayons » (combustion), confirmé par les deux documents de
# référence — déjà promis en Phase 1 par CLAUDE.md, jamais implémenté avant
# ce jalon.
COMBUSTION_ORB_DEGREES = 15.0

# Version courte de la dignité essentielle pour les phrases relationnelles
# (Ascendant/maître, Luminaires...) — distincte de DIGNITY_LABELS ci-dessus
# (utilisée dans les puces "Dignités et réceptions", qui gardent la
# parenthèse "(détriment)" et le préfixe "En").
DIGNITY_INLINE_CLAUSE = {
    "Domicile": "en domicile",
    "Exaltation": "en exaltation",
    "Exil (détriment)": "en exil",
    "Chute": "en chute",
    "Pérégrin": "pérégrin",
    "Pérégrine": "pérégrine",
}


def format_dms(decimal_degrees: float) -> str:
    """Formate un degré décimal dans le signe en notation degrés/minutes.

    Arrondit et reporte la retenue sur le degré plutôt que de produire une
    notation invalide comme "27°60'" (60' = 1°).
    """
    degrees = int(decimal_degrees)
    minutes = round((decimal_degrees - degrees) * 60)
    if minutes == 60:
        minutes = 0
        degrees += 1
    return f"{degrees}°{minutes:02d}'"


def direction_label(name: str, retrograde: bool | None) -> str:
    if retrograde is None:
        return "—"
    if retrograde:
        return "Rétrograde"
    return "Directe" if name in FEMININE_PLANETS else "Direct"


def _positions_row(table, point: PointPosition) -> None:
    cells = table.add_row().cells
    cells[0].text = point.name
    cells[1].text = point.sign
    cells[2].text = format_dms(point.degree_in_sign)
    cells[3].text = str(point.house)
    cells[4].text = point.sect_role or "—"
    cells[5].text = direction_label(point.name, point.retrograde)
    cells[6].text = point.essential_dignity or "—"


def add_positions_table(document: Document, observation: Observation):
    table = document.add_table(rows=1, cols=len(POSITIONS_HEADER))
    for cell, text in zip(table.rows[0].cells, POSITIONS_HEADER):
        styles.set_header_cell(cell, text)
    styles.style_table(table, POSITIONS_COLUMN_WIDTHS_DXA)

    for point in observation.all_points:
        _positions_row(table, point)
    return table


def add_rulerships_table(document: Document, observation: Observation):
    table = document.add_table(rows=1, cols=len(RULERSHIPS_HEADER))
    for cell, text in zip(table.rows[0].cells, RULERSHIPS_HEADER):
        styles.set_header_cell(cell, text)
    styles.style_table(table, RULERSHIPS_COLUMN_WIDTHS_DXA)

    for rulership in observation.rulerships:
        cells = table.add_row().cells
        cells[0].text = rulership.planet
        cells[1].text = ", ".join(rulership.domicile_signs)
        cells[2].text = ", ".join(str(house) for house in rulership.houses_governed)
    return table


def add_minor_dignities_table(document: Document, observation: Observation):
    """Table des dignités mineures (triplicité, bornes égyptiennes, décans) —
    séparée de la table des positions, qui correspond exactement au tableau
    des documents de référence (lesquels ne couvrent pas ces dignités)."""
    table = document.add_table(rows=1, cols=len(MINOR_DIGNITIES_HEADER))
    for cell, text in zip(table.rows[0].cells, MINOR_DIGNITIES_HEADER):
        styles.set_header_cell(cell, text)
    styles.style_table(table, MINOR_DIGNITIES_COLUMN_WIDTHS_DXA)

    for planet in observation.planets:
        cells = table.add_row().cells
        cells[0].text = planet.name
        cells[1].text = planet.triplicity_dignity or "—"
        cells[2].text = planet.bound_dignity or "—"
        cells[3].text = planet.decan_dignity or "—"
    return table


def format_releasing_date(when: datetime) -> str:
    return when.strftime("%d/%m/%Y")


def _zodiacal_releasing_row(
    table, period: ReleasingPeriod, level_label: str, fortune_sign: str
) -> None:
    cells = table.add_row().cells
    cells[0].text = level_label
    cells[1].text = period.sign
    cells[2].text = period.ruler
    cells[3].text = format_releasing_date(period.start)
    cells[4].text = format_releasing_date(period.end)
    # Toujours par rapport à la Part de Fortune, même dans la table de
    # l'Esprit (confirmé par une source primaire, voir zodiacal_releasing.is_peak_period).
    cells[5].text = "Oui" if is_peak_period(period, fortune_sign) else "—"


def add_zodiacal_releasing_table(
    document: Document, chapters: list[ReleasingChapter], fortune_sign: str
):
    """Table des chapitres (L1) et sous-périodes (L2) de libération
    zodiacale. `fortune_sign` sert uniquement au marquage des périodes
    culminantes ; les périodes elles-mêmes sont déjà calculées (Fortune ou
    Esprit) en amont dans `core.chart`."""
    table = document.add_table(rows=1, cols=len(ZODIACAL_RELEASING_HEADER))
    for cell, text in zip(table.rows[0].cells, ZODIACAL_RELEASING_HEADER):
        styles.set_header_cell(cell, text)
    styles.style_table(table, ZODIACAL_RELEASING_COLUMN_WIDTHS_DXA)

    for chapter in chapters:
        _zodiacal_releasing_row(table, chapter.l1, "L1", fortune_sign)
        for sub in chapter.sub_periods:
            _zodiacal_releasing_row(table, sub, "L2", fortune_sign)
    return table


def _display_name(name: str) -> str:
    return DISPLAY_NAME_WITH_ARTICLE.get(name, name)


def _join_french(names: list[str]) -> str:
    if len(names) == 1:
        return names[0]
    return ", ".join(names[:-1]) + " et " + names[-1]


def conjunction_text(cluster: SignCluster) -> str | None:
    """Puce de conjonction pour un amas d'au moins deux membres.

    Renvoie None pour un amas à un seul membre (rien à conjoindre).
    """
    if len(cluster.members) < 2:
        return None
    names = _join_french([_display_name(m) for m in cluster.members])
    verb = "conjointes" if all(m in FEMININE_POINTS for m in cluster.members) else "conjoints"
    return f"{names} {verb} en {cluster.sign} (maison {cluster.house})."


def cluster_aspect_text(cluster_aspect: ClusterAspect, clusters_by_sign: dict[str, SignCluster]) -> str:
    cluster_a = clusters_by_sign[cluster_aspect.sign_a]
    cluster_b = clusters_by_sign[cluster_aspect.sign_b]
    if cluster_aspect.boundary_exception:
        return (
            f"{cluster_a.sign} (maison {cluster_a.house}) et {cluster_b.sign} "
            f"(maison {cluster_b.house}) : conjonction hors signe "
            "(règle des 3°, signes adjacents)."
        )
    label = ASPECT_LABEL[cluster_aspect.aspect]
    return (
        f"{cluster_a.sign} (maison {cluster_a.house}) en {label} avec "
        f"{cluster_b.sign} (maison {cluster_b.house})."
    )


def mutual_reception_text(reception: MutualReception) -> str:
    return f"Réception mutuelle par domicile entre {reception.planet_a} et {reception.planet_b}."


def add_aspects_section(document: Document, observation: Observation) -> None:
    """Puces factuelles des aspects par signe : conjonctions intra-amas,
    relation d'aspect (ou aversion) pour chaque paire d'amas, puis les
    réceptions mutuelles par domicile (un lien technique indépendant des
    aspects, placé ici comme dans les documents de référence).

    Le commentaire interprétatif (ex. significations de maîtrise) reste hors
    périmètre de docgen : c'est une tâche de rédaction, pas de calcul.
    """
    for cluster in observation.clusters:
        text = conjunction_text(cluster)
        if text:
            document.add_paragraph(text, style="List Bullet")

    clusters_by_sign = {cluster.sign: cluster for cluster in observation.clusters}
    for cluster_aspect in observation.cluster_aspects:
        document.add_paragraph(
            cluster_aspect_text(cluster_aspect, clusters_by_sign), style="List Bullet"
        )

    for reception in observation.mutual_receptions:
        document.add_paragraph(mutual_reception_text(reception), style="List Bullet")


def _distribution_points(observation: Observation) -> list[PointPosition]:
    """Le jeu de points utilisé pour la répartition élémentaire/modale et
    l'angularité (Phase 2) : Ascendant + 7 planètes + Nœud Nord + les 3
    Parts, à l'exclusion du Milieu du Ciel et du Nœud Sud. Convention déduite
    par inspection croisée des deux documents de référence, qui l'appliquent
    tous les deux identiquement sans jamais l'expliciter."""
    excluded = {"Milieu du Ciel", "Nœud Sud"}
    return [p for p in observation.all_points if p.name not in excluded]


def add_elemental_modal_section(document: Document, observation: Observation) -> None:
    """Répartition élémentaire (Ascendant/Soleil/Lune) et modale (les 12
    points de `_distribution_points`), en Phase 2.

    Rendu volontairement plus simple, en style « libellé technique », que la
    prose des documents de référence (qui emploie un jugement qualitatif —
    « dominante fixe », « suivie de près par » — ne généralisant pas
    proprement à un comptage arbitraire, et une accord grammatical
    singulier/pluriel qui ajouterait de la fragilité pour peu de gain) :
    même choix que le jalon 4 pour les puces d'aspects.
    """
    points = _distribution_points(observation)
    main_points = [observation.ascendant, observation.planet("Soleil"), observation.planet("Lune")]
    main_names = {"Ascendant", "Soleil", "Lune"}

    # Liste à virgules simples (pas de "et" final) : convention des documents
    # de référence pour ce type d'énumération positionnelle, distincte de la
    # jonction "et" utilisée par _join_french pour les puces d'aspects.
    main_elements_text = ", ".join(
        f"{p.name} en {p.element.lower()} ({p.sign})" for p in main_points
    )
    present_among_main = {p.element for p in main_points}
    missing_elements = sorted({"Feu", "Terre", "Air", "Eau"} - present_among_main)
    other_points = [p for p in points if p.name not in main_names]
    elements_elsewhere = {p.element for p in other_points}
    present_elsewhere = [e for e in missing_elements if e in elements_elsewhere]
    absent_entirely = [e for e in missing_elements if e not in elements_elsewhere]

    def _lower_list(elements: list[str]) -> str:
        return _join_french([e.lower() for e in elements]) if elements else "aucun"

    document.add_paragraph(
        f"Éléments des trois points directeurs : {main_elements_text}. "
        f"Éléments absents de ces trois points : {_lower_list(missing_elements)}. "
        f"Présents ailleurs dans le thème : {_lower_list(present_elsewhere)}. "
        f"Absents de l'ensemble du thème : {_lower_list(absent_entirely)}."
    )

    modality_groups: dict[str, list[PointPosition]] = {"Cardinal": [], "Fixe": [], "Mutable": []}
    for p in points:
        modality_groups[p.modality].append(p)
    ordered_modalities = sorted(modality_groups.items(), key=lambda kv: -len(kv[1]))

    modality_sentences = []
    for modality, members in ordered_modalities:
        count = len(members)
        names = _join_french([p.name for p in members]) if members else "aucun facteur"
        modality_sentences.append(f"{modality} : {names} ({count} facteur{'s' if count != 1 else ''}).")
    document.add_paragraph(
        "Modalité, par ordre décroissant de nombre de facteurs : " + " ".join(modality_sentences)
    )


def add_angularity_section(document: Document, observation: Observation) -> None:
    """Angularité (maisons angulaires 1/4/7/10) des 12 points de
    `_distribution_points`, en Phase 2. Structure identique dans les deux
    documents de référence, reproduite fidèlement — à l'exception du nom
    complet « le Soleil » employé une fois dans le document de Liam, non
    repris ici : incohérent avec la convention déjà établie et testée en
    Phase 1 (`DISPLAY_NAME_WITH_ARTICLE`), qui n'ajoute pas d'article aux
    luminaires/planètes.
    """
    points = _distribution_points(observation)
    by_house: dict[int, list[PointPosition]] = {}
    for p in points:
        by_house.setdefault(p.house, []).append(p)

    angular_sentences = [
        f"Maison {house} : {', '.join(p.name for p in by_house[house])}."
        for house in (1, 4, 7, 10)
        if by_house.get(house)
    ]
    document.add_paragraph(" ".join(angular_sentences))

    non_angular_houses = [h for h in by_house if h not in (1, 4, 7, 10)]
    if non_angular_houses:
        groups = [
            f"{_join_french([_display_name(p.name) for p in by_house[house]])} "
            f"(maison {house}, {house_quality(house).lower()})"
            for house in non_angular_houses
        ]
        document.add_paragraph("Hors angularité : " + ", ".join(groups) + ".")


def add_dignities_and_receptions_section(document: Document, observation: Observation) -> None:
    """Dignités et réceptions, en Phase 2 : liste à puces (comme les aspects,
    jalon 4), pas de rédaction relationnelle — réutilise presque entièrement
    des données déjà calculées (dignité essentielle, réception mutuelle,
    rétrogradation), plus la combustion (« sous les rayons »).

    Simplifications assumées par rapport aux documents de référence, qui
    divergent stylistiquement sur ces points (même logique que le jalon 4) :
    pas de qualificatif subjectif "juste" pour la planète non combuste la
    plus proche du seuil ; titre "Rétrogrades" invariant (pas d'accord
    singulier/pluriel ni de qualificatif "(parmi les sept planètes...)").
    """
    by_category: dict[str, list[PointPosition]] = {}
    peregrines = []
    for planet in observation.planets:
        dignity = planet.essential_dignity
        if dignity in ("Pérégrin", "Pérégrine"):
            peregrines.append(planet)
        elif dignity is not None:
            by_category.setdefault(dignity, []).append(planet)

    # Listes à virgules simples (pas de "et" final), y compris pour 2
    # éléments : convention confirmée par les deux documents de référence
    # pour les catégories de dignité et les pérégrins spécifiquement,
    # distincte de la jonction _join_french utilisée plus bas (réceptions,
    # combustion, rétrogrades) qui, elle, ajoute "et" avant le dernier élément.
    for category in DIGNITY_CATEGORY_ORDER:
        members = by_category.get(category)
        if members:
            names = ", ".join(f"{p.name} ({p.sign})" for p in members)
            document.add_paragraph(f"{DIGNITY_LABELS[category]} : {names}.", style="List Bullet")

    if peregrines:
        names = ", ".join(p.name for p in peregrines)
        document.add_paragraph(
            f"Pérégrins (sans dignité essentielle) : {names}.", style="List Bullet"
        )

    for reception in observation.mutual_receptions:
        document.add_paragraph(
            f"Réception mutuelle par domicile : {reception.planet_a} et {reception.planet_b}.",
            style="List Bullet",
        )

    combust = [sp for sp in observation.solar_proximity if sp.gap_degrees < COMBUSTION_ORB_DEGREES]
    non_combust = [
        sp for sp in observation.solar_proximity if sp.gap_degrees >= COMBUSTION_ORB_DEGREES
    ]
    if combust:
        combust_text = _join_french(
            [f"{sp.planet} ({format_dms(sp.gap_degrees)} d'écart)" for sp in combust]
        )
        sentence = f"Sous les rayons du Soleil (moins de {COMBUSTION_ORB_DEGREES:g}°) : {combust_text}."
        if non_combust:
            nearest = min(non_combust, key=lambda sp: sp.gap_degrees)
            sentence += (
                f" {nearest.planet} est hors de cette configuration "
                f"({format_dms(nearest.gap_degrees)} d'écart)."
            )
        document.add_paragraph(sentence, style="List Bullet")

    retrogrades = [p.name for p in observation.planets if p.retrograde]
    if retrogrades:
        document.add_paragraph(f"Rétrogrades : {_join_french(retrogrades)}.", style="List Bullet")


def _conjunction_clause(point: PointPosition, observation: Observation) -> str | None:
    """« conjoint(e) à X, Y » pour les autres membres de l'amas de `point`
    (lui-même exclu). None si `point` est seul dans son signe.

    Réutilisable pour toute sous-section relationnelle (Ascendant/maître,
    Luminaires, Nœuds et Parts) : un point conjoint à ses co-résidents de
    signe s'exprime toujours de la même façon.
    """
    cluster = next(c for c in observation.clusters if c.sign == point.sign)
    others = [m for m in cluster.members if m != point.name]
    if not others:
        return None
    verb = "conjointe" if point.name in FEMININE_POINTS else "conjoint"
    return f"{verb} à {_join_french([_display_name(m) for m in others])}"


def _rulership_aspect_clause(ruler: PointPosition, governed_sign: str) -> str | None:
    """Relation entre le signe de `ruler` et `governed_sign` (le signe qu'il
    gouverne) — None si `ruler` est déjà dans ce signe (cas couvert par
    `_conjunction_clause`, une aspect à soi-même n'aurait pas de sens).

    Pas d'article devant `governed_sign` (simplification assumée : accorder
    le genre/nombre des noms de signes en français — Gémeaux/Poissons
    pluriels, Vierge/Balance féminins — ajouterait une table de genre inutile
    ; cohérent avec `cluster_aspect_text`, Phase 1, qui n'en met pas non
    plus). Le pronom "qu'il/qu'elle gouverne" s'accorde au genre du maître
    (la planète, déjà dans FEMININE_PLANETS), pas du signe.
    """
    if ruler.sign == governed_sign:
        return None
    pronoun = "elle" if ruler.name in FEMININE_PLANETS else "il"
    aspect = sign_aspect(ruler.sign, governed_sign)
    return f"en {ASPECT_LABEL[aspect]} avec {governed_sign} qu'{pronoun} gouverne"


def add_ascendant_and_ruler_section(document: Document, observation: Observation) -> None:
    """Ascendant et son maître, en Phase 2 : première sous-section à décrire
    une *relation* (pas seulement lister des faits) — voir `_conjunction_clause`
    et `_rulership_aspect_clause`, conçues pour être réutilisées par les
    sous-sections Luminaires et Nœuds et Parts à venir.
    """
    ascendant = observation.ascendant
    ruler_name = next(planet for planet, signs in DOMICILES.items() if ascendant.sign in signs)
    ruler = observation.planet(ruler_name)
    rulership = next(r for r in observation.rulerships if r.planet == ruler_name)

    ascendant_clause = _conjunction_clause(ascendant, observation) or "sans planète en maison 1"
    document.add_paragraph(f"Ascendant en {ascendant.sign}, maison 1, {ascendant_clause}.")

    regent_clause = (
        "seul régent du" if len(rulership.domicile_signs) == 1 else "l'un des deux régents traditionnels du"
    )
    clauses = [DIGNITY_INLINE_CLAUSE[ruler.essential_dignity]]
    ruler_conjunction = _conjunction_clause(ruler, observation)
    if ruler_conjunction:
        clauses.append(ruler_conjunction)
    rulership_aspect = _rulership_aspect_clause(ruler, ascendant.sign)
    if rulership_aspect:
        clauses.append(rulership_aspect)
    sentence = (
        f"Maître de l'Ascendant : {ruler.name}, {regent_clause} {ascendant.sign}, "
        f"situé en {ruler.sign}, maison {ruler.house}, {', '.join(clauses)}."
    )
    document.add_paragraph(sentence)

    if len(rulership.domicile_signs) > 1:
        other_house = next(
            house
            for sign, house in zip(rulership.domicile_signs, rulership.houses_governed)
            if sign != ascendant.sign
        )
        document.add_paragraph(f"{ruler.name} régit également la maison {other_house}.")


def _cluster_display_name(cluster: SignCluster) -> str:
    """Nom d'un amas cible dans une clause relationnelle : "l'amas du Signe"
    s'il a plusieurs membres (convention déjà établie par `cluster_aspect_text`,
    Phase 1), le nom de son unique membre sinon."""
    if len(cluster.members) == 1:
        return _display_name(cluster.members[0])
    return f"l'amas du {cluster.sign}"


def _cluster_relations_clause(point: PointPosition, observation: Observation) -> str | None:
    """Relations d'aspect de l'amas de `point` vers tous les autres amas,
    groupées par type d'aspect (y compris aversion) dans l'ordre canonique de
    `ASPECT_LABEL`, chaque groupe joint par `_join_french`. None si l'amas de
    `point` n'a de relation avec aucun autre (cas non rencontré en pratique).

    Réservée à la Lune dans `add_luminaries_section` : les deux documents de
    référence énumèrent toujours ces relations pour la Lune, jamais pour le
    Soleil (remplacé par `_solar_rays_clause`).
    """
    clusters_by_sign = {c.sign: c for c in observation.clusters}
    targets_by_aspect: dict[str, list[SignCluster]] = {}
    for cluster_aspect in observation.cluster_aspects:
        if cluster_aspect.sign_a == point.sign:
            other = clusters_by_sign[cluster_aspect.sign_b]
        elif cluster_aspect.sign_b == point.sign:
            other = clusters_by_sign[cluster_aspect.sign_a]
        else:
            continue
        targets_by_aspect.setdefault(cluster_aspect.aspect, []).append(other)

    if not targets_by_aspect:
        return None

    clauses = []
    for aspect, label in ASPECT_LABEL.items():
        targets = targets_by_aspect.get(aspect)
        if not targets:
            continue
        ordered = sorted(targets, key=lambda c: index_of_sign(c.sign))
        names = _join_french([_cluster_display_name(c) for c in ordered])
        clauses.append(f"en {label} avec {names}")
    return ", ".join(clauses)


def _solar_rays_clause(observation: Observation) -> str:
    """Clause de combustion du point de vue du Soleil lui-même : "sous les
    rayons de X" pour les planètes combustes (réutilise `solar_proximity`,
    déjà calculé pour "Dignités et réceptions", jalon 20), ou la clause fixe
    des deux documents de référence si aucune ne l'est."""
    combust = [sp for sp in observation.solar_proximity if sp.gap_degrees < COMBUSTION_ORB_DEGREES]
    if not combust:
        return "sous les rayons de personne (il est la source)"
    names = _join_french([f"{sp.planet} ({format_dms(sp.gap_degrees)} d'écart)" for sp in combust])
    return f"sous les rayons de {names}"


def _sect_light_clause(point: PointPosition, observation: Observation) -> str | None:
    """"lumière de secte de ce thème diurne/nocturne" pour le luminaire en
    accord avec la secte de la carte, None sinon. Un seul des deux documents
    de référence l'explicite (Anthony, pour sa Lune nocturne) ; retenue de
    façon symétrique pour les deux luminaires (même logique que la règle de
    secte elle-même, voir CLAUDE.md)."""
    if point.sect_role != "Lumière de secte":
        return None
    return f"lumière de secte de ce thème {observation.sect}"


def add_luminaries_section(document: Document, observation: Observation) -> None:
    """Luminaires (Soleil, Lune), en Phase 2 : premier texte de Phase 2 où
    Soleil et Lune reçoivent des traitements non symétriques, confirmé par
    les deux documents de référence — la Lune énumère toujours ses relations
    d'aspect à tous les autres amas (`_cluster_relations_clause`), le Soleil
    ne le fait jamais, remplacé par la clause de combustion de son propre
    point de vue (`_solar_rays_clause`). Inclut la mention "lumière de secte"
    (`_sect_light_clause`) pour le luminaire en accord avec la secte de la
    carte, et la phase de lunaison natale (`core.lunation`, voir ce module
    pour une divergence relevée mais non reproduite sur le thème de Liam).
    """
    for point in (observation.planet("Soleil"), observation.planet("Lune")):
        rulership = next(r for r in observation.rulerships if r.planet == point.name)
        clauses = [DIGNITY_INLINE_CLAUSE[point.essential_dignity]]

        sect_clause = _sect_light_clause(point, observation)
        if sect_clause:
            clauses.append(sect_clause)

        if point.name == "Soleil":
            clauses.append(_solar_rays_clause(observation))

        conjunction = _conjunction_clause(point, observation)
        if conjunction:
            clauses.append(conjunction)

        if point.name == "Lune":
            relations = _cluster_relations_clause(point, observation)
            if relations:
                clauses.append(relations)

        document.add_paragraph(
            f"{point.name} : {point.sign}, maison {point.house}, "
            f"régit la maison {rulership.houses_governed[0]}, {', '.join(clauses)}."
        )

    phase = observation.lunation_phase
    document.add_paragraph(
        f"Phase de lunaison natale : {phase.name} "
        f"(écart Soleil-Lune d'environ {round(phase.gap_degrees)}°)."
    )


def build_observation_document(observation: Observation) -> Document:
    """Construit le document .docx : Phase 1 (Observation) complète, et les
    sous-sections de Phase 2 (Fiche technique) déjà couvertes par docgen
    (voir CLAUDE.md pour l'état d'avancement exact).

    Phase 3 (Interprétation) et le reste de la Phase 2 impliquent une
    rédaction descriptive qui dépasse encore le périmètre de docgen : docgen
    ne fait que mettre en forme des données déjà calculées.
    """
    document = Document()
    document.add_heading("Phase 1 — Observation", level=1)

    document.add_heading("Positions planétaires et facteurs sensibles", level=2)
    add_positions_table(document, observation)

    document.add_heading(
        "Maîtrises traditionnelles (règle de rulership unique par planète)", level=2
    )
    add_rulerships_table(document, observation)

    document.add_heading("Dignités mineures (triplicité, bornes, décans)", level=2)
    add_minor_dignities_table(document, observation)

    document.add_heading("Aspects par signe relevés", level=2)
    add_aspects_section(document, observation)

    fortune_sign = observation.part_of_fortune.sign
    document.add_heading("Libération zodiacale — Part de Fortune", level=2)
    add_zodiacal_releasing_table(document, observation.zodiacal_releasing_fortune, fortune_sign)

    document.add_heading("Libération zodiacale — Part de l'Esprit", level=2)
    add_zodiacal_releasing_table(document, observation.zodiacal_releasing_spirit, fortune_sign)

    document.add_heading("Phase 2 — Fiche technique", level=1)

    document.add_heading("Répartition élémentaire et modale", level=2)
    add_elemental_modal_section(document, observation)

    document.add_heading("Angularité", level=2)
    add_angularity_section(document, observation)

    document.add_heading("Dignités et réceptions", level=2)
    add_dignities_and_receptions_section(document, observation)

    document.add_heading("Ascendant et son maître", level=2)
    add_ascendant_and_ruler_section(document, observation)

    document.add_heading("Luminaires", level=2)
    add_luminaries_section(document, observation)

    return document
