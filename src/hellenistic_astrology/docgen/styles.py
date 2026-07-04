"""Mise en forme des tableaux `python-docx` (en-têtes, largeurs de
colonnes, ombrage) — utilisé par `builder.py`, purement visuel."""

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips

HEADER_SHADING = "EDE6F0"
FONT_NAME = "Calibri"
HEADER_FONT_SIZE_PT = 10

# Ombrage des cellules de dignité essentielle dans la table des positions
# (jalon 33) : vert pour Domicile/Exaltation (dignités favorables), rouge
# pour Exil/Chute (défavorables). Pérégrin reste sans ombrage (neutre).
DIGNITY_FAVORABLE_SHADING = "D9EAD3"
DIGNITY_UNFAVORABLE_SHADING = "F4CCCC"


def shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_header_cell(cell, text: str) -> None:
    cell.text = text
    shade_cell(cell, HEADER_SHADING)
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(HEADER_FONT_SIZE_PT)


def style_table(table, column_widths_dxa: list[int]) -> None:
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, column_widths_dxa):
            cell.width = Twips(width)


def add_table_of_contents(document) -> None:
    """Insère un champ Word "Table des matières" (`TOC \\o "1-2" \\h \\z \\u`,
    titres de niveau 1-2, hyperliens, sans numéro de page pour les entrées
    masquées). python-docx n'a pas de support natif pour les tables des
    matières — approche standard documentée : construire directement les
    éléments oxml bas niveau (`w:fldChar` begin/separate/end +
    `w:instrText`). Limite assumée : la pagination réelle est calculée par
    le moteur de mise en page de Word, pas par ce projet — le champ reste
    vide (placeholder "Mettre à jour le champ...") tant que Word ne l'a pas
    peuplé, ce que `_set_update_fields_on_open` en dessous déclenche
    automatiquement à l'ouverture du document."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Mettre à jour le champ pour afficher la table des matières."

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(instr_text)
    run._r.append(field_separate)
    run._r.append(placeholder)
    run._r.append(field_end)

    _set_update_fields_on_open(document)


def _set_update_fields_on_open(document) -> None:
    """Demande à Word de recalculer les champs (dont la TOC ci-dessus) à
    l'ouverture du document, plutôt que de laisser l'utilisateur faire un
    clic-droit -> Mettre à jour les champs manuellement."""
    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
