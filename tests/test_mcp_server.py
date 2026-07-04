import json

import docx
import pytest

from hellenistic_astrology.mcp_server import (
    _assemble_final_document,
    _compute_observation_json,
    _generate_document,
    _generate_interpretation_brief,
    _with_error_handling,
)

from .regression_helpers import load_fixture


@pytest.mark.parametrize(
    "fixture_name, expected_ascendant_sign",
    [("anthony", "Lion"), ("liam", "Verseau")],
)
def test_compute_observation_json_matches_known_fixture(fixture_name, expected_ascendant_sign):
    birth_data = load_fixture(fixture_name)["birth_data"]

    result = _compute_observation_json(birth_data)
    parsed = json.loads(result)

    assert parsed["ascendant"]["sign"] == expected_ascendant_sign
    assert parsed["name"] == birth_data["name"]


def test_generate_document_writes_readable_docx(tmp_path):
    birth_data = load_fixture("anthony")["birth_data"]
    target = tmp_path / "anthony.docx"

    result = _generate_document(birth_data, str(target))

    assert target.exists()
    assert str(target) in result
    document = docx.Document(str(target))
    assert any(p.text == "Phase 1 — Observation" for p in document.paragraphs)


def test_generate_document_default_path_uses_slugified_name(tmp_path, monkeypatch):
    import hellenistic_astrology.mcp_server as mcp_server_module

    monkeypatch.setattr(mcp_server_module, "OUTPUT_DIR", tmp_path)
    birth_data = load_fixture("liam")["birth_data"]

    result = _generate_document(birth_data)

    expected = tmp_path / "liam.docx"
    assert expected.exists()
    assert str(expected) in result


def test_generate_interpretation_brief_contains_expected_sections(tmp_path, monkeypatch):
    import hellenistic_astrology.mcp_server as mcp_server_module

    monkeypatch.setattr(mcp_server_module, "OUTPUT_DIR", tmp_path)
    birth_data = load_fixture("anthony")["birth_data"]

    brief = _generate_interpretation_brief(birth_data)

    assert "## Orientation générale" in brief
    assert "## Repères temporels actuels" in brief
    assert (tmp_path / "anthony_brief.md").read_text(encoding="utf-8") == brief


def test_assemble_final_document_appends_and_returns_confirmation(tmp_path):
    docx_path = tmp_path / "anthony.docx"
    markdown_path = tmp_path / "anthony_phase3_draft.md"
    docx.Document().save(docx_path)
    markdown_path.write_text("# Phase 3 — Interprétation\n\nUn paragraphe.\n", encoding="utf-8")

    result = _assemble_final_document(str(docx_path), str(markdown_path))

    expected = tmp_path / "anthony_final.docx"
    assert result == f"écrit : {expected}"
    assert expected.exists()


def test_assemble_final_document_missing_file_returns_error_message(tmp_path):
    result = _with_error_handling(_assemble_final_document, str(tmp_path / "absent.docx"), str(tmp_path / "absent.md"))

    assert result.startswith("erreur : ")


def test_with_error_handling_returns_message_on_value_error():
    def _boom(birth_data):
        raise ValueError("données de naissance incomplètes")

    result = _with_error_handling(_boom, {"name": "Test"})

    assert result == "erreur : données de naissance incomplètes"


def test_compute_observation_json_missing_birth_fields_returns_error_message():
    result = _with_error_handling(_compute_observation_json, {"name": "Test sans données"})

    assert result.startswith("erreur : ")
