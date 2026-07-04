import docx
import pytest

from hellenistic_astrology.assembly import (
    MarkdownBlock,
    append_markdown_to_document,
    assemble_final_document,
    parse_markdown_blocks,
)


def test_parse_markdown_blocks_headings_paragraph_and_bullet():
    text = """# Phase 3 — Interprétation

Un paragraphe normal.

## Orientation générale

Un autre paragraphe.

## Bibliographie

- Auteur, *Titre en italique* — https://example.com
"""
    blocks = parse_markdown_blocks(text)

    assert blocks[0] == MarkdownBlock(kind="heading1", segments=[("Phase 3 — Interprétation", False)])
    assert blocks[1] == MarkdownBlock(kind="paragraph", segments=[("Un paragraphe normal.", False)])
    assert blocks[2] == MarkdownBlock(kind="heading2", segments=[("Orientation générale", False)])
    assert blocks[3] == MarkdownBlock(kind="paragraph", segments=[("Un autre paragraphe.", False)])
    assert blocks[4] == MarkdownBlock(kind="heading2", segments=[("Bibliographie", False)])
    assert blocks[5] == MarkdownBlock(
        kind="bullet",
        segments=[("Auteur, ", False), ("Titre en italique", True), (" — https://example.com", False)],
    )


def test_parse_markdown_blocks_skips_blank_lines():
    blocks = parse_markdown_blocks("# Titre\n\n\n\nParagraphe.\n")
    assert len(blocks) == 2


def test_append_markdown_to_document_produces_expected_styles_and_runs():
    document = docx.Document()

    append_markdown_to_document(document, "# Titre\n## Sous-titre\nTexte *important* ici.\n- Une puce\n")

    paragraphs = document.paragraphs
    assert [p.style.name for p in paragraphs] == ["Heading 1", "Heading 2", "Normal", "List Bullet"]
    assert paragraphs[0].text == "Titre"
    assert paragraphs[1].text == "Sous-titre"
    assert paragraphs[2].text == "Texte important ici."
    runs = paragraphs[2].runs
    assert [r.text for r in runs] == ["Texte ", "important", " ici."]
    assert [bool(r.italic) for r in runs] == [False, True, False]
    assert paragraphs[3].text == "Une puce"


def test_assemble_final_document_appends_and_does_not_touch_input(tmp_path):
    docx_path = tmp_path / "anthony.docx"
    markdown_path = tmp_path / "anthony_phase3_draft.md"

    base_document = docx.Document()
    base_document.add_heading("Phase 1 — Observation", level=1)
    base_document.save(docx_path)
    original_bytes = docx_path.read_bytes()

    markdown_path.write_text("# Phase 3 — Interprétation\n\nUn paragraphe rédigé.\n", encoding="utf-8")

    result = assemble_final_document(docx_path, markdown_path)

    assert result == tmp_path / "anthony_final.docx"
    assert result.exists()
    # Le fichier d'entrée n'est jamais modifié.
    assert docx_path.read_bytes() == original_bytes

    final_document = docx.Document(result)
    texts = [p.text for p in final_document.paragraphs]
    assert texts == ["Phase 1 — Observation", "Phase 3 — Interprétation", "Un paragraphe rédigé."]


def test_assemble_final_document_custom_output_path(tmp_path):
    docx_path = tmp_path / "anthony.docx"
    markdown_path = tmp_path / "brouillon.md"
    output_path = tmp_path / "sortie" / "complet.docx"

    docx.Document().save(docx_path)
    markdown_path.write_text("# Titre\n", encoding="utf-8")

    result = assemble_final_document(docx_path, markdown_path, output_path)

    assert result == output_path
    assert output_path.exists()


def test_assemble_final_document_missing_docx_raises_value_error(tmp_path):
    markdown_path = tmp_path / "brouillon.md"
    markdown_path.write_text("# Titre\n", encoding="utf-8")

    with pytest.raises(ValueError, match="introuvable"):
        assemble_final_document(tmp_path / "absent.docx", markdown_path)


def test_assemble_final_document_missing_markdown_raises_value_error(tmp_path):
    docx_path = tmp_path / "anthony.docx"
    docx.Document().save(docx_path)

    with pytest.raises(ValueError, match="introuvable"):
        assemble_final_document(docx_path, tmp_path / "absent.md")
