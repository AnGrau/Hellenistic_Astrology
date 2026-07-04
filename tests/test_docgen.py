from datetime import datetime

from docx import Document
import pytest

from hellenistic_astrology.core import dignities as dignities_module
from hellenistic_astrology.core.aspects import ClusterAspect, SignCluster, sign_aspect
from hellenistic_astrology.core.chart import build_observation
from hellenistic_astrology.core.dignities import MutualReception, Rulership, SolarProximity
from hellenistic_astrology.core.eclipse import EclipseConfiguration
from hellenistic_astrology.core.lunation import LunationPhase
from hellenistic_astrology.core.observation import Observation, PointPosition
from hellenistic_astrology.core.zodiacal_releasing import ReleasingChapter, ReleasingPeriod
from hellenistic_astrology.docgen.builder import (
    ASPECT_GLYPH,
    MINOR_DIGNITIES_HEADER,
    POSITIONS_HEADER,
    RULERSHIPS_HEADER,
    ZODIACAL_RELEASING_HEADER,
    add_angularity_section,
    add_ascendant_and_ruler_section,
    add_aspectarian_table,
    add_cover_page,
    add_dignities_and_receptions_section,
    add_elemental_modal_section,
    add_luminaries_section,
    add_minor_dignities_table,
    add_nodes_and_parts_section,
    add_positions_table,
    add_zodiacal_releasing_table,
    build_observation_document,
    cluster_aspect_text,
    conjunction_text,
    direction_label,
    format_dms,
    format_releasing_date,
    mutual_reception_text,
)
from hellenistic_astrology.docgen import styles

from .regression_helpers import birth_data_from_fixture, load_fixture


@pytest.mark.parametrize(
    "decimal_degrees, expected",
    [
        (28.1833, "28°11'"),
        (0.2167, "0°13'"),
        (19.9833, "19°59'"),
        (27.9999, "28°00'"),  # ne doit jamais produire "27°60'"
    ],
)
def test_format_dms(decimal_degrees, expected):
    assert format_dms(decimal_degrees) == expected


def test_direction_label_gender_agreement():
    assert direction_label("Soleil", False) == "Direct"
    assert direction_label("Lune", False) == "Directe"
    assert direction_label("Vénus", True) == "Rétrograde"
    assert direction_label("Mars", True) == "Rétrograde"
    assert direction_label("Ascendant", None) == "—"


def test_conjunction_text_single_member_is_none():
    cluster = SignCluster(sign="Balance", house=3, members=("Mars",))
    assert conjunction_text(cluster) is None


def test_conjunction_text_mixed_gender_uses_masculine_plural():
    cluster = SignCluster(sign="Lion", house=1, members=("Ascendant", "Lune"))
    assert conjunction_text(cluster) == "l'Ascendant et Lune conjoints en Lion (maison 1)."


def test_conjunction_text_all_feminine_uses_feminine_plural():
    cluster = SignCluster(sign="Scorpion", house=4, members=("Lune", "Vénus"))
    assert conjunction_text(cluster) == "Lune et Vénus conjointes en Scorpion (maison 4)."


def test_conjunction_text_three_members_with_article():
    cluster = SignCluster(
        sign="Taureau", house=10, members=("Saturne", "Milieu du Ciel", "Part de l'Esprit")
    )
    assert conjunction_text(cluster) == (
        "Saturne, le Milieu du Ciel et la Part de l'Esprit conjoints en Taureau (maison 10)."
    )


def test_cluster_aspect_text_real_aspect():
    clusters_by_sign = {
        "Lion": SignCluster(sign="Lion", house=1, members=("Ascendant",)),
        "Scorpion": SignCluster(sign="Scorpion", house=4, members=("Soleil",)),
    }
    aspect = ClusterAspect(sign_a="Lion", sign_b="Scorpion", aspect="Carré")
    assert cluster_aspect_text(aspect, clusters_by_sign) == (
        "Lion (maison 1) en carré avec Scorpion (maison 4)."
    )


def test_cluster_aspect_text_boundary_exception():
    clusters_by_sign = {
        "Scorpion": SignCluster(sign="Scorpion", house=4, members=("Lune",)),
        "Sagittaire": SignCluster(sign="Sagittaire", house=5, members=("Saturne",)),
    }
    aspect = ClusterAspect(
        sign_a="Scorpion", sign_b="Sagittaire", aspect="Aversion", boundary_exception=True
    )
    assert cluster_aspect_text(aspect, clusters_by_sign) == (
        "Scorpion (maison 4) et Sagittaire (maison 5) : conjonction hors signe "
        "(règle des 3°, signes adjacents)."
    )


def test_mutual_reception_text():
    reception = MutualReception(planet_a="Vénus", planet_b="Mars")
    assert mutual_reception_text(reception) == (
        "Réception mutuelle par domicile entre Vénus et Mars."
    )


def test_add_minor_dignities_table_uses_dash_for_missing_values():
    ascendant = PointPosition(name="Ascendant", sign="Bélier", degree_in_sign=0, house=1)
    midheaven = PointPosition(name="Milieu du Ciel", sign="Capricorne", degree_in_sign=0, house=10)
    planets = [
        PointPosition(
            name="Mars",
            sign="Bélier",
            degree_in_sign=3,
            house=1,
            triplicity_dignity="Maître de triplicité (jour)",
            bound_dignity=None,
            decan_dignity="Maître du décan",
        ),
        PointPosition(name="Vénus", sign="Bélier", degree_in_sign=3, house=1),
    ]
    observation = Observation(
        name="Test", sect="diurne", ascendant=ascendant, midheaven=midheaven, planets=planets
    )
    document = Document()

    table = add_minor_dignities_table(document, observation)

    mars_row = [cell.text for cell in table.rows[1].cells]
    assert mars_row == ["Mars", "Maître de triplicité (jour)", "—", "Maître du décan"]
    venus_row = [cell.text for cell in table.rows[2].cells]
    assert venus_row == ["Vénus", "—", "—", "—"]


def test_format_releasing_date():
    assert format_releasing_date(datetime(1970, 11, 20)) == "20/11/1970"


def test_add_zodiacal_releasing_table_rows_and_peak_marking():
    # Fortune synthétique en Scorpion : Lion (angulaire, groupe fixe) doit
    # être marqué culminant, Vierge (non angulaire à Scorpion) ne doit pas
    # l'être. Vierge est aussi marquée comme un relâchement du lien pour
    # tester le rendu de la nouvelle colonne (Lion, non issu d'un saut, sert
    # de cas négatif).
    fortune_sign = "Scorpion"
    l1 = ReleasingPeriod(
        level=1, sign="Lion", ruler="Soleil", start=datetime(2000, 1, 1), end=datetime(2019, 1, 1)
    )
    l2 = ReleasingPeriod(
        level=2, sign="Vierge", ruler="Mercure", start=datetime(2000, 1, 1), end=datetime(2001, 8, 1),
        bond_loosed=True,
    )
    chapters = [ReleasingChapter(l1=l1, sub_periods=[l2])]
    document = Document()

    table = add_zodiacal_releasing_table(document, chapters, fortune_sign)

    header_row = [cell.text for cell in table.rows[0].cells]
    assert header_row == ZODIACAL_RELEASING_HEADER
    l1_row = [cell.text for cell in table.rows[1].cells]
    assert l1_row == ["L1", "Lion", "Soleil", "01/01/2000", "01/01/2019", "Oui", "—"]
    l2_row = [cell.text for cell in table.rows[2].cells]
    assert l2_row == ["L2", "Vierge", "Mercure", "01/01/2000", "01/08/2001", "—", "Oui"]


def test_add_elemental_modal_and_angularity_sections_synthetic():
    # Nœud Sud et Milieu du Ciel présents dans all_points mais doivent être
    # exclus des deux sections (convention des 12 points, voir CLAUDE.md
    # jalon 19) : inclus ici justement pour vérifier qu'ils sont bien ignorés.
    ascendant = PointPosition(
        name="Ascendant", sign="Bélier", degree_in_sign=0, house=1,
        element="Feu", modality="Cardinal", house_quality="Angulaire",
    )
    soleil = PointPosition(
        name="Soleil", sign="Lion", degree_in_sign=0, house=4,
        element="Feu", modality="Fixe", house_quality="Angulaire",
    )
    lune = PointPosition(
        name="Lune", sign="Cancer", degree_in_sign=0, house=1,
        element="Eau", modality="Cardinal", house_quality="Angulaire",
    )
    mercure = PointPosition(
        name="Mercure", sign="Vierge", degree_in_sign=0, house=2,
        element="Terre", modality="Mutable", house_quality="Succédente",
    )
    north_node = PointPosition(
        name="Nœud Nord", sign="Balance", degree_in_sign=0, house=3,
        element="Air", modality="Cardinal", house_quality="Cadente",
    )
    south_node = PointPosition(
        name="Nœud Sud", sign="Bélier", degree_in_sign=0, house=1,
        element="Feu", modality="Cardinal", house_quality="Angulaire",
    )
    midheaven = PointPosition(
        name="Milieu du Ciel", sign="Capricorne", degree_in_sign=0, house=10,
        element="Terre", modality="Cardinal", house_quality="Angulaire",
    )
    fortune = PointPosition(
        name="Part de Fortune", sign="Scorpion", degree_in_sign=0, house=3,
        element="Eau", modality="Fixe", house_quality="Cadente",
    )
    observation = Observation(
        name="Test",
        sect="diurne",
        ascendant=ascendant,
        midheaven=midheaven,
        planets=[soleil, lune, mercure],
        part_of_fortune=fortune,
        north_node=north_node,
        south_node=south_node,
        all_points=[ascendant, soleil, lune, mercure, north_node, south_node, midheaven, fortune],
    )
    document = Document()

    add_elemental_modal_section(document, observation)
    add_angularity_section(document, observation)

    paragraphs = [p.text for p in document.paragraphs]
    assert paragraphs[0] == (
        "Éléments des trois points directeurs : Ascendant en feu (Bélier), "
        "Soleil en feu (Lion), Lune en eau (Cancer). "
        "Éléments absents de ces trois points : air et terre. "
        "Présents ailleurs dans le thème : air et terre. "
        "Absents de l'ensemble du thème : aucun."
    )
    assert paragraphs[1] == (
        "Modalité, par ordre décroissant de nombre de facteurs : "
        "Cardinal : Ascendant, Lune et Nœud Nord (3 facteurs). "
        "Fixe : Soleil et Part de Fortune (2 facteurs). "
        "Mutable : Mercure (1 facteur)."
    )
    # paragraphs[2] est le paragraphe du graphique élément/modalité inséré
    # par add_elemental_modal_section (jalon 33) : pas de texte, une image.
    assert paragraphs[3] == "Maison 1 : Ascendant, Lune. Maison 4 : Soleil."
    assert paragraphs[4] == (
        "Hors angularité : Mercure (maison 2, succédente), "
        "Nœud Nord et la Part de Fortune (maison 3, cadente)."
    )


def test_add_dignities_and_receptions_section_synthetic():
    def make_planet(name, sign, essential_dignity, retrograde):
        return PointPosition(
            name=name, sign=sign, degree_in_sign=0, house=1,
            essential_dignity=essential_dignity, retrograde=retrograde,
        )

    planets = [
        make_planet("Soleil", "Balance", "Chute", False),
        make_planet("Lune", "Taureau", "Exaltation", False),
        make_planet("Mercure", "Gémeaux", "Pérégrin", False),
        make_planet("Vénus", "Balance", "Pérégrine", True),
        make_planet("Mars", "Taureau", "Exil (détriment)", True),
        make_planet("Jupiter", "Sagittaire", "Pérégrin", False),
        make_planet("Saturne", "Lion", "Exil (détriment)", False),
    ]
    ascendant = PointPosition(name="Ascendant", sign="Bélier", degree_in_sign=0, house=1)
    midheaven = PointPosition(name="Milieu du Ciel", sign="Capricorne", degree_in_sign=0, house=10)
    observation = Observation(
        name="Test",
        sect="diurne",
        ascendant=ascendant,
        midheaven=midheaven,
        planets=planets,
        mutual_receptions=[MutualReception(planet_a="Mars", planet_b="Vénus")],
        solar_proximity=[
            SolarProximity(planet="Mercure", gap_degrees=10.0),
            SolarProximity(planet="Vénus", gap_degrees=20.0),
            SolarProximity(planet="Mars", gap_degrees=25.0),
            SolarProximity(planet="Jupiter", gap_degrees=5.0),
            SolarProximity(planet="Saturne", gap_degrees=40.0),
        ],
    )
    document = Document()

    add_dignities_and_receptions_section(document, observation)

    bullets = [p.text for p in document.paragraphs if p.style.name == "List Bullet"]
    assert bullets == [
        "En chute : Soleil (Balance).",
        "En exaltation : Lune (Taureau).",
        "En exil (détriment) : Mars (Taureau), Saturne (Lion).",
        "Pérégrins (sans dignité essentielle) : Mercure, Vénus, Jupiter.",
        "Réception mutuelle par domicile : Mars et Vénus.",
        "Sous les rayons du Soleil (moins de 15°) : Mercure (10°00' d'écart) et Jupiter (5°00' d'écart). "
        "Vénus est hors de cette configuration (20°00' d'écart).",
        "Rétrogrades : Vénus et Mars.",
    ]


def test_add_ascendant_and_ruler_section_single_domicile_with_aspect():
    # Bélier n'est pas utilisé ici : maître (Soleil) à domicile unique, avec
    # aspect (carré) entre son signe et celui de l'Ascendant, et conjonction
    # des deux côtés — configuration proche d'Anthony (jalon 21).
    ascendant = PointPosition(name="Ascendant", sign="Lion", degree_in_sign=0, house=1)
    lune = PointPosition(name="Lune", sign="Lion", degree_in_sign=0, house=1)
    soleil = PointPosition(
        name="Soleil", sign="Scorpion", degree_in_sign=0, house=4, essential_dignity="Pérégrin"
    )
    venus = PointPosition(name="Vénus", sign="Scorpion", degree_in_sign=0, house=4)
    observation = Observation(
        name="Test",
        sect="diurne",
        ascendant=ascendant,
        midheaven=PointPosition(name="Milieu du Ciel", sign="Taureau", degree_in_sign=0, house=10),
        planets=[soleil, lune, venus],
        rulerships=[Rulership(planet="Soleil", domicile_signs=("Lion",), houses_governed=(1,))],
        clusters=[
            SignCluster(sign="Lion", house=1, members=("Ascendant", "Lune")),
            SignCluster(sign="Scorpion", house=4, members=("Soleil", "Vénus")),
        ],
    )
    document = Document()

    add_ascendant_and_ruler_section(document, observation)

    paragraphs = [p.text for p in document.paragraphs]
    assert paragraphs[0] == "Ascendant en Lion, maison 1, conjoint à Lune."
    assert paragraphs[1] == (
        "Maître de l'Ascendant : Soleil, seul régent du Lion, situé en Scorpion, maison 4, "
        "pérégrin, conjoint à Vénus, en carré avec Lion qu'il gouverne."
    )
    assert len(paragraphs) == 2  # pas de "régit également" : domicile unique.


def test_add_ascendant_and_ruler_section_double_domicile_aversion_alone():
    # Verseau : maître (Saturne) à double domicile, seul dans son signe (pas
    # de conjonction), en aversion avec le signe qu'il gouverne — configuration
    # proche de Liam pour la structure, mais aversion plutôt qu'opposition
    # (cas non illustré par les documents de référence, à couvrir ici).
    ascendant = PointPosition(name="Ascendant", sign="Verseau", degree_in_sign=0, house=1)
    saturne = PointPosition(
        name="Saturne", sign="Poissons", degree_in_sign=0, house=2, essential_dignity="Pérégrin"
    )
    observation = Observation(
        name="Test",
        sect="diurne",
        ascendant=ascendant,
        midheaven=PointPosition(name="Milieu du Ciel", sign="Scorpion", degree_in_sign=0, house=10),
        planets=[saturne],
        rulerships=[
            Rulership(planet="Saturne", domicile_signs=("Capricorne", "Verseau"), houses_governed=(12, 1))
        ],
        clusters=[
            SignCluster(sign="Verseau", house=1, members=("Ascendant",)),
            SignCluster(sign="Poissons", house=2, members=("Saturne",)),
        ],
    )
    document = Document()

    add_ascendant_and_ruler_section(document, observation)

    paragraphs = [p.text for p in document.paragraphs]
    assert paragraphs[0] == "Ascendant en Verseau, maison 1, sans planète en maison 1."
    assert paragraphs[1] == (
        "Maître de l'Ascendant : Saturne, l'un des deux régents traditionnels du Verseau, "
        "situé en Poissons, maison 2, pérégrin, en aversion avec Verseau qu'il gouverne."
    )
    assert paragraphs[2] == "Saturne régit également la maison 12."


def test_add_ascendant_and_ruler_section_ruler_in_ascendant_sign():
    # Le maître est dans le signe de l'Ascendant lui-même : pas de clause
    # d'aspect (déjà couverte par la conjonction), mais "régit également"
    # reste présent (double domicile).
    ascendant = PointPosition(name="Ascendant", sign="Bélier", degree_in_sign=0, house=1)
    mars = PointPosition(
        name="Mars", sign="Bélier", degree_in_sign=0, house=1, essential_dignity="Domicile"
    )
    observation = Observation(
        name="Test",
        sect="diurne",
        ascendant=ascendant,
        midheaven=PointPosition(name="Milieu du Ciel", sign="Capricorne", degree_in_sign=0, house=10),
        planets=[mars],
        rulerships=[
            Rulership(planet="Mars", domicile_signs=("Bélier", "Scorpion"), houses_governed=(1, 8))
        ],
        clusters=[SignCluster(sign="Bélier", house=1, members=("Ascendant", "Mars"))],
    )
    document = Document()

    add_ascendant_and_ruler_section(document, observation)

    paragraphs = [p.text for p in document.paragraphs]
    assert paragraphs[0] == "Ascendant en Bélier, maison 1, conjoint à Mars."
    assert paragraphs[1] == (
        "Maître de l'Ascendant : Mars, l'un des deux régents traditionnels du Bélier, "
        "situé en Bélier, maison 1, en domicile, conjoint à l'Ascendant."
    )
    assert paragraphs[2] == "Mars régit également la maison 8."


def test_add_luminaries_section_synthetic():
    # Configuration synthétique couvrant les deux traitements non symétriques :
    # Soleil (pas de relation d'aspect énumérée, combustion de son propre
    # point de vue, ici vide -> clause fixe) vs Lune (relations d'aspect à
    # tous les autres amas énumérées, groupées par type). Soleil est lumière
    # de secte (thème diurne), la Lune ne l'est pas -> teste les deux branches
    # de `_sect_light_clause`.
    ascendant = PointPosition(name="Ascendant", sign="Bélier", degree_in_sign=0, house=1)
    soleil = PointPosition(
        name="Soleil", sign="Lion", degree_in_sign=0, house=5,
        essential_dignity="Domicile", sect_role="Lumière de secte",
    )
    mercure = PointPosition(name="Mercure", sign="Lion", degree_in_sign=0, house=5)
    lune = PointPosition(
        name="Lune", sign="Cancer", degree_in_sign=0, house=8,
        essential_dignity="Domicile", sect_role="Hors secte (jour)",
    )
    observation = Observation(
        name="Test",
        sect="diurne",
        ascendant=ascendant,
        midheaven=PointPosition(name="Milieu du Ciel", sign="Capricorne", degree_in_sign=0, house=10),
        planets=[soleil, mercure, lune],
        rulerships=[
            Rulership(planet="Soleil", domicile_signs=("Lion",), houses_governed=(5,)),
            Rulership(planet="Lune", domicile_signs=("Cancer",), houses_governed=(8,)),
        ],
        clusters=[
            SignCluster(sign="Bélier", house=1, members=("Ascendant",)),
            SignCluster(sign="Cancer", house=8, members=("Lune",)),
            SignCluster(sign="Lion", house=5, members=("Soleil", "Mercure")),
            SignCluster(sign="Verseau", house=11, members=("Saturne", "Vénus")),
        ],
        cluster_aspects=[
            ClusterAspect(sign_a="Bélier", sign_b="Cancer", aspect="Carré"),
            ClusterAspect(sign_a="Cancer", sign_b="Verseau", aspect="Aversion"),
            ClusterAspect(sign_a="Bélier", sign_b="Lion", aspect="Trigone"),
            ClusterAspect(sign_a="Lion", sign_b="Verseau", aspect="Opposition"),
        ],
        solar_proximity=[],
        lunation_phase=LunationPhase(name="gibbeuse", gap_degrees=150.4),
    )
    document = Document()

    add_luminaries_section(document, observation)

    paragraphs = [p.text for p in document.paragraphs]
    assert paragraphs == [
        "Soleil : Lion, maison 5, régit la maison 5, en domicile, "
        "lumière de secte de ce thème diurne, sous les rayons de personne (il est la source), "
        "conjoint à Mercure.",
        "Lune : Cancer, maison 8, régit la maison 8, en domicile, "
        "en carré avec l'Ascendant, en aversion avec l'amas du Verseau.",
        "Phase de lunaison natale : gibbeuse (écart Soleil-Lune d'environ 150°).",
    ]


def test_add_luminaries_section_sun_combustion_and_no_sect_light():
    # Soleil hors secte (thème nocturne) et combuste par deux planètes :
    # teste la clause de combustion non vide, jamais couverte par le test
    # synthétique précédent, ainsi que l'absence de clause de secte.
    soleil = PointPosition(
        name="Soleil", sign="Scorpion", degree_in_sign=0, house=4,
        essential_dignity="Pérégrin", sect_role="Hors secte (nuit)",
    )
    lune = PointPosition(
        name="Lune", sign="Lion", degree_in_sign=0, house=1,
        essential_dignity="Pérégrine", sect_role="Lumière de secte",
    )
    observation = Observation(
        name="Test",
        sect="nocturne",
        ascendant=PointPosition(name="Ascendant", sign="Lion", degree_in_sign=0, house=1),
        midheaven=PointPosition(name="Milieu du Ciel", sign="Taureau", degree_in_sign=0, house=10),
        planets=[soleil, lune],
        rulerships=[
            Rulership(planet="Soleil", domicile_signs=("Lion",), houses_governed=(1,)),
            Rulership(planet="Lune", domicile_signs=("Cancer",), houses_governed=(12,)),
        ],
        clusters=[
            SignCluster(sign="Lion", house=1, members=("Lune",)),
            SignCluster(sign="Scorpion", house=4, members=("Soleil",)),
        ],
        cluster_aspects=[ClusterAspect(sign_a="Lion", sign_b="Scorpion", aspect="Carré")],
        solar_proximity=[
            SolarProximity(planet="Mercure", gap_degrees=13.5),
            SolarProximity(planet="Jupiter", gap_degrees=9.0),
            SolarProximity(planet="Vénus", gap_degrees=20.0),
        ],
        lunation_phase=LunationPhase(name="disséminatrice", gap_degrees=269.5),
    )
    document = Document()

    add_luminaries_section(document, observation)

    paragraphs = [p.text for p in document.paragraphs]
    assert paragraphs[0] == (
        "Soleil : Scorpion, maison 4, régit la maison 1, pérégrin, "
        "sous les rayons de Mercure (13°30' d'écart) et Jupiter (9°00' d'écart)."
    )
    assert paragraphs[1] == (
        "Lune : Lion, maison 1, régit la maison 12, pérégrine, "
        "lumière de secte de ce thème nocturne, en carré avec Soleil."
    )


def test_add_nodes_and_parts_section_synthetic_with_conjunctions_and_eclipse():
    # Les deux Nœuds ont chacun un co-résident de signe (teste
    # `_conjunction_clause` pour les deux), l'écart Lune-Nœud est sous le
    # seuil solaire (mention de proximité), la relation Fortune/Esprit est un
    # vrai aspect (trigone, pas une aversion), et la configuration est une
    # éclipse solaire (branche positive de `_eclipse_configuration_clause`,
    # non illustrée par les deux thèmes de référence qui n'en ont aucune).
    ascendant = PointPosition(name="Ascendant", sign="Cancer", degree_in_sign=0, house=1)
    venus = PointPosition(name="Vénus", sign="Bélier", degree_in_sign=0, house=3)
    soleil = PointPosition(name="Soleil", sign="Balance", degree_in_sign=0, house=9)
    north_node = PointPosition(name="Nœud Nord", sign="Bélier", degree_in_sign=0, house=3)
    south_node = PointPosition(name="Nœud Sud", sign="Balance", degree_in_sign=0, house=9)
    fortune = PointPosition(name="Part de Fortune", sign="Lion", degree_in_sign=0, house=7)
    spirit = PointPosition(name="Part de l'Esprit", sign="Sagittaire", degree_in_sign=0, house=11)
    observation = Observation(
        name="Test",
        sect="diurne",
        ascendant=ascendant,
        midheaven=PointPosition(name="Milieu du Ciel", sign="Verseau", degree_in_sign=0, house=8),
        planets=[venus, soleil],
        part_of_fortune=fortune,
        part_of_spirit=spirit,
        north_node=north_node,
        south_node=south_node,
        clusters=[
            SignCluster(sign="Bélier", house=3, members=("Nœud Nord", "Vénus")),
            SignCluster(sign="Balance", house=9, members=("Nœud Sud", "Soleil")),
            SignCluster(sign="Lion", house=7, members=("Part de Fortune",)),
            SignCluster(sign="Sagittaire", house=11, members=("Part de l'Esprit",)),
        ],
        eclipse=EclipseConfiguration(
            is_eclipse=True, eclipse_type="solaire", node_gap_degrees=5.0, closer_node="Nœud Nord",
            syzygy_type="Nouvelle Lune", syzygy_gap_degrees=3.0,
        ),
    )
    document = Document()

    add_nodes_and_parts_section(document, observation)

    paragraphs = [p.text for p in document.paragraphs]
    assert paragraphs == [
        "Axe des Nœuds : Nœud Nord en Bélier (maison 3), conjoint à Vénus, "
        "Nœud Sud en Balance (maison 9), conjoint à Soleil, "
        "la Lune est proche du Nœud Nord en degrés réels (5°00' d'écart).",
        "La Part de Fortune (Lion, maison 7) en trigone avec la Part de l'Esprit (Sagittaire, maison 11).",
        "Configuration d'éclipse à la naissance (éclipse solaire) : écart Lune-Nœud de 5°00', "
        "écart Soleil-Lune de 3°00' par rapport à la Nouvelle Lune.",
    ]


def test_add_nodes_and_parts_section_synthetic_no_conjunction_no_eclipse():
    # Les deux Nœuds sont seuls dans leur signe (pas de clause de
    # conjonction), l'écart Lune-Nœud dépasse le seuil solaire (pas de
    # mention de proximité), et la relation Fortune/Esprit est une aversion.
    ascendant = PointPosition(name="Ascendant", sign="Poissons", degree_in_sign=0, house=1)
    north_node = PointPosition(name="Nœud Nord", sign="Cancer", degree_in_sign=0, house=4)
    south_node = PointPosition(name="Nœud Sud", sign="Capricorne", degree_in_sign=0, house=10)
    fortune = PointPosition(name="Part de Fortune", sign="Taureau", degree_in_sign=0, house=2)
    spirit = PointPosition(name="Part de l'Esprit", sign="Gémeaux", degree_in_sign=0, house=3)
    observation = Observation(
        name="Test",
        sect="nocturne",
        ascendant=ascendant,
        midheaven=PointPosition(name="Milieu du Ciel", sign="Sagittaire", degree_in_sign=0, house=10),
        planets=[],
        part_of_fortune=fortune,
        part_of_spirit=spirit,
        north_node=north_node,
        south_node=south_node,
        clusters=[
            SignCluster(sign="Taureau", house=2, members=("Part de Fortune",)),
            SignCluster(sign="Gémeaux", house=3, members=("Part de l'Esprit",)),
            SignCluster(sign="Cancer", house=4, members=("Nœud Nord",)),
            SignCluster(sign="Capricorne", house=10, members=("Nœud Sud",)),
        ],
        eclipse=EclipseConfiguration(
            is_eclipse=False, eclipse_type=None, node_gap_degrees=25.0, closer_node="Nœud Sud",
            syzygy_type="Pleine Lune", syzygy_gap_degrees=40.0,
        ),
    )
    document = Document()

    add_nodes_and_parts_section(document, observation)

    paragraphs = [p.text for p in document.paragraphs]
    assert paragraphs == [
        "Axe des Nœuds : Nœud Nord en Cancer (maison 4), Nœud Sud en Capricorne (maison 10).",
        "La Part de Fortune (Taureau, maison 2) et la Part de l'Esprit (Gémeaux, maison 3) sont en aversion : "
        "aucun aspect ptoléméen ne les relie directement.",
        "Aucune configuration d'éclipse à la naissance : écart Lune-Nœud de 25°00', "
        "écart Soleil-Lune de 40°00' par rapport à la Pleine Lune.",
    ]


def _cell_shading_fill(cell) -> str | None:
    shd = cell._tc.get_or_add_tcPr().find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
    )
    if shd is None:
        return None
    return shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")


def test_add_positions_table_shades_dignity_cells():
    def make_planet(name, sign, essential_dignity):
        return PointPosition(
            name=name, sign=sign, degree_in_sign=0, house=1, essential_dignity=essential_dignity,
        )

    ascendant = PointPosition(name="Ascendant", sign="Bélier", degree_in_sign=0, house=1)
    midheaven = PointPosition(name="Milieu du Ciel", sign="Capricorne", degree_in_sign=0, house=10)
    domicile_planet = make_planet("Mars", "Bélier", "Domicile")
    exil_planet = make_planet("Vénus", "Bélier", "Exil (détriment)")
    peregrine_planet = make_planet("Saturne", "Gémeaux", "Pérégrin")
    observation = Observation(
        name="Test", sect="diurne", ascendant=ascendant, midheaven=midheaven,
        planets=[domicile_planet, exil_planet, peregrine_planet],
        all_points=[ascendant, domicile_planet, exil_planet, peregrine_planet, midheaven],
    )
    document = Document()

    table = add_positions_table(document, observation)

    rows_by_name = {row.cells[0].text: row for row in table.rows[1:]}
    assert _cell_shading_fill(rows_by_name["Mars"].cells[6]) == styles.DIGNITY_FAVORABLE_SHADING
    assert _cell_shading_fill(rows_by_name["Vénus"].cells[6]) == styles.DIGNITY_UNFAVORABLE_SHADING
    assert _cell_shading_fill(rows_by_name["Saturne"].cells[6]) is None
    assert _cell_shading_fill(rows_by_name["Ascendant"].cells[6]) is None


def test_add_aspectarian_table_triangular_structure():
    ascendant = PointPosition(name="Ascendant", sign="Bélier", degree_in_sign=0, house=1)
    midheaven = PointPosition(name="Milieu du Ciel", sign="Capricorne", degree_in_sign=0, house=10)
    planets = [
        PointPosition(name=name, sign=sign, degree_in_sign=0, house=1)
        for name, sign in [
            ("Soleil", "Bélier"), ("Lune", "Bélier"), ("Mercure", "Cancer"),
            ("Vénus", "Balance"), ("Mars", "Capricorne"), ("Jupiter", "Lion"),
            ("Saturne", "Verseau"),
        ]
    ]
    observation = Observation(
        name="Test", sect="diurne", ascendant=ascendant, midheaven=midheaven, planets=planets,
        all_points=[ascendant, *planets, midheaven],
    )
    document = Document()

    table = add_aspectarian_table(document, observation)

    assert [cell.text for cell in table.rows[0].cells] == [
        "", "Soleil", "Lune", "Mercure", "Vénus", "Mars", "Jupiter",
    ]
    assert [row.cells[0].text for row in table.rows[1:]] == [
        "Lune", "Mercure", "Vénus", "Mars", "Jupiter", "Saturne",
    ]
    # Même signe (Soleil/Lune, tous deux Bélier) : glyphe de conjonction.
    assert table.rows[1].cells[1].text == "☌"
    # Case triangulaire supérieure jamais renseignée (ex. Lune/Lune n'existe
    # pas, mais Mercure/Vénus, au-dessus de la diagonale, doit rester vide).
    assert table.rows[2].cells[4].text == ""
    # Dernière ligne (Saturne) : toutes les 6 colonnes renseignées.
    assert all(table.rows[6].cells[i].text for i in range(1, 7))


def test_add_cover_page_centers_name_ascendant_sect_and_wheel():
    ascendant = PointPosition(name="Ascendant", sign="Lion", degree_in_sign=10, house=1)
    midheaven = PointPosition(name="Milieu du Ciel", sign="Taureau", degree_in_sign=5, house=10)
    soleil = PointPosition(name="Soleil", sign="Scorpion", degree_in_sign=10, house=4)
    observation = Observation(
        name="Anthony", sect="nocturne", ascendant=ascendant, midheaven=midheaven,
        planets=[soleil], all_points=[ascendant, soleil, midheaven],
    )
    document = Document()

    add_cover_page(document, observation)

    assert document.paragraphs[0].text == "Anthony"
    assert document.paragraphs[0].style.name == "Title"
    assert document.paragraphs[1].text == "Ascendant Lion (maison 1) — Thème nocturne"
    assert len(document.inline_shapes) == 1


def test_add_table_of_contents_inserts_field_and_update_setting():
    document = Document()

    styles.add_table_of_contents(document)

    paragraph_xml = document.paragraphs[-1]._p.xml
    assert "fldChar" in paragraph_xml
    assert "instrText" in paragraph_xml
    assert 'TOC \\o "1-2" \\h \\z \\u' in paragraph_xml
    assert "updateFields" in document.settings.element.xml


@pytest.mark.parametrize("fixture_name", ["anthony", "liam"])
def test_build_observation_document_structure(fixture_name):
    fixture = load_fixture(fixture_name)
    observation = build_observation(birth_data_from_fixture(fixture))

    document = build_observation_document(observation)

    assert [p.text for p in document.paragraphs if p.style.name == "Heading 1"] == [
        "Table des matières",
        "Phase 1 — Observation",
        "Phase 2 — Fiche technique",
    ]
    assert len(document.tables) == 6

    (
        positions_table,
        rulerships_table,
        minor_dignities_table,
        aspectarian_table,
        zr_fortune_table,
        zr_spirit_table,
    ) = document.tables

    header_row = [cell.text for cell in positions_table.rows[0].cells]
    assert header_row == POSITIONS_HEADER
    # Ascendant + 7 planètes classiques + Nœud Nord + Nœud Sud + MC + Fortune + Esprit + Éros.
    assert len(positions_table.rows) == 1 + 1 + 7 + 2 + 1 + 3

    ascendant_row = [cell.text for cell in positions_table.rows[1].cells]
    assert ascendant_row[0] == "Ascendant"
    assert ascendant_row[1] == fixture["ascendant"]["sign"]
    assert ascendant_row[3] == str(fixture["ascendant"]["house"])

    sun_row = next(
        [cell.text for cell in row.cells]
        for row in positions_table.rows
        if row.cells[0].text == "Soleil"
    )
    expected_sun = fixture["planets"]["Soleil"]
    assert sun_row[1] == expected_sun["sign"]
    assert sun_row[3] == str(expected_sun["house"])
    assert sun_row[4] == expected_sun["sect_role"]
    assert sun_row[6] == expected_sun["essential_dignity"]

    ruler_header = [cell.text for cell in rulerships_table.rows[0].cells]
    assert ruler_header == RULERSHIPS_HEADER
    assert len(rulerships_table.rows) == 1 + 7

    mercury_row = next(
        [cell.text for cell in row.cells]
        for row in rulerships_table.rows
        if row.cells[0].text == "Mercure"
    )
    expected_mercury = fixture["rulerships"]["Mercure"]
    assert mercury_row[1] == ", ".join(expected_mercury["domicile_signs"])
    assert mercury_row[2] == ", ".join(str(h) for h in expected_mercury["houses_governed"])

    minor_header = [cell.text for cell in minor_dignities_table.rows[0].cells]
    assert minor_header == MINOR_DIGNITIES_HEADER
    assert len(minor_dignities_table.rows) == 1 + 7

    # Recalcule les valeurs attendues depuis les mêmes fonctions core.dignities,
    # appliquées à la position réellement calculée (pas la fixture, sujette à
    # une tolérance d'arrondi qui pourrait tomber près d'une frontière de
    # borne/décan) : ceci teste le branchement docgen, pas l'exactitude
    # astrologique déjà couverte par tests/test_dignities.py.
    for row in minor_dignities_table.rows[1:]:
        planet_name, triplicity_cell, bound_cell, decan_cell = (c.text for c in row.cells)
        actual_planet = observation.planet(planet_name)
        assert triplicity_cell == (
            dignities_module.triplicity_dignity(planet_name, actual_planet.sign) or "—"
        )
        assert bound_cell == (
            dignities_module.bound_dignity(
                planet_name, actual_planet.sign, actual_planet.degree_in_sign
            )
            or "—"
        )
        assert decan_cell == (
            dignities_module.decan_dignity(
                planet_name, actual_planet.sign, actual_planet.degree_in_sign
            )
            or "—"
        )

    heading2_texts = [p.text for p in document.paragraphs if p.style.name == "Heading 2"]
    assert heading2_texts[-6:] == [
        "Répartition élémentaire et modale",
        "Angularité",
        "Dignités et réceptions",
        "Ascendant et son maître",
        "Luminaires",
        "Nœuds et Parts",
    ]
    assert "Dignités mineures (triplicité, bornes, décans)" in heading2_texts
    assert "Aspects par signe relevés" in heading2_texts
    assert "Aspectarian (planète × planète)" in heading2_texts
    assert "Libération zodiacale — Part de Fortune" in heading2_texts
    assert "Libération zodiacale — Part de l'Esprit" in heading2_texts

    # Aspectarian : recoupe une case au hasard (Soleil/Lune) contre le même
    # calcul core.aspects.sign_aspect, pas un nouveau calcul.
    assert [cell.text for cell in aspectarian_table.rows[0].cells][0] == ""
    soleil_sign = observation.planet("Soleil").sign
    lune_sign = observation.planet("Lune").sign
    lune_row = next(
        row for row in aspectarian_table.rows if row.cells[0].text == "Lune"
    )
    expected_aspect = sign_aspect(soleil_sign, lune_sign)
    expected_glyph = "☌" if expected_aspect is None else ASPECT_GLYPH[expected_aspect]
    assert lune_row.cells[1].text == expected_glyph

    # Recoupe le rendu docgen contre le même calcul core.zodiacal_releasing,
    # sur la position réellement calculée (pas la fixture, qui ne documente
    # pas cette technique) : teste le branchement, pas l'exactitude
    # astrologique déjà couverte par tests/test_zodiacal_releasing.py.
    for table, chapters in [
        (zr_fortune_table, observation.zodiacal_releasing_fortune),
        (zr_spirit_table, observation.zodiacal_releasing_spirit),
    ]:
        assert [cell.text for cell in table.rows[0].cells] == ZODIACAL_RELEASING_HEADER
        expected_row_count = 1 + sum(1 + len(chapter.sub_periods) for chapter in chapters)
        assert len(table.rows) == expected_row_count
        first_l1_row = [cell.text for cell in table.rows[1].cells]
        assert first_l1_row[0] == "L1"
        assert first_l1_row[1] == chapters[0].l1.sign
        assert first_l1_row[3] == format_releasing_date(chapters[0].l1.start)

    fixture_clusters = [
        SignCluster(sign=c["sign"], house=c["house"], members=tuple(c["members"]))
        for c in fixture["clusters"]
    ]
    clusters_by_sign = {c.sign: c for c in fixture_clusters}
    expected_bullets = [
        text for text in (conjunction_text(c) for c in fixture_clusters) if text is not None
    ]
    expected_bullets += [
        cluster_aspect_text(
            ClusterAspect(
                sign_a=a["sign_a"],
                sign_b=a["sign_b"],
                aspect=a["aspect"],
                boundary_exception=a["boundary_exception"],
            ),
            clusters_by_sign,
        )
        for a in fixture["cluster_aspects"]
    ]
    expected_bullets += [
        mutual_reception_text(MutualReception(planet_a=r["planet_a"], planet_b=r["planet_b"]))
        for r in fixture["mutual_receptions"]
    ]

    # Recoupe la section Dignités et réceptions en l'appelant directement sur
    # la même observation (déjà testée en isolation par le test synthétique
    # ci-dessus) : vérifie ici seulement le branchement dans
    # build_observation_document, pas l'exactitude du gabarit lui-même.
    dignities_document = Document()
    add_dignities_and_receptions_section(dignities_document, observation)
    expected_bullets += [
        p.text for p in dignities_document.paragraphs if p.style.name == "List Bullet"
    ]

    actual_bullets = [p.text for p in document.paragraphs if p.style.name == "List Bullet"]
    assert actual_bullets == expected_bullets

    # La phrase "Maison N : ..." est reproduite mot pour mot contre le texte
    # réel des documents de référence (seule sous-section de Phase 2 dont la
    # structure est validée identique dans les deux documents ; Anthony omet
    # le Milieu du Ciel de la maison 10 ici, contrairement au document
    # d'origine — voir CLAUDE.md jalon 19 pour la justification de cette
    # divergence assumée). La phrase "Hors angularité" n'est *pas* reproduite
    # mot pour mot (simplifications assumées, voir CLAUDE.md).
    expected_angular_sentence = {
        "anthony": (
            "Maison 1 : Ascendant, Lune. Maison 4 : Soleil, Vénus, Jupiter, Part de Fortune. "
            "Maison 7 : Nœud Nord, Part d'Éros. Maison 10 : Saturne, Part de l'Esprit."
        ),
        "liam": (
            "Maison 1 : Ascendant. Maison 4 : Lune, Mars. Maison 7 : Saturne, Part de Fortune. "
            "Maison 10 : Mercure."
        ),
    }
    all_paragraphs = list(document.paragraphs)
    angularity_index = next(
        i
        for i, p in enumerate(all_paragraphs)
        if p.style.name == "Heading 2" and p.text == "Angularité"
    )
    angularity_paragraph = next(p for p in all_paragraphs[angularity_index + 1 :] if p.text)
    assert angularity_paragraph.text == expected_angular_sentence[fixture_name]

    # "Ascendant et son maître" recoupé mot pour mot contre le texte réel des
    # deux documents, à l'ordre signe/maison près (retenu : signe puis maison,
    # voir CLAUDE.md jalon 21) et à la relation maître↔Ascendant près pour
    # Anthony (omise dans le document d'origine, toujours incluse ici).
    expected_ascendant_paragraphs = {
        "anthony": [
            "Ascendant en Lion, maison 1, conjoint à Lune et Nœud Sud.",
            "Maître de l'Ascendant : Soleil, seul régent du Lion, situé en Scorpion, maison 4, "
            "pérégrin, conjoint à Vénus, Jupiter et la Part de Fortune, en carré avec Lion qu'il gouverne.",
        ],
        "liam": [
            "Ascendant en Verseau, maison 1, sans planète en maison 1.",
            "Maître de l'Ascendant : Saturne, l'un des deux régents traditionnels du Verseau, "
            "situé en Lion, maison 7, en exil, conjoint à la Part de Fortune, "
            "en opposition avec Verseau qu'il gouverne.",
            "Saturne régit également la maison 12.",
        ],
    }
    ascendant_heading_index = next(
        i
        for i, p in enumerate(all_paragraphs)
        if p.style.name == "Heading 2" and p.text == "Ascendant et son maître"
    )
    luminaries_heading_index = next(
        i
        for i, p in enumerate(all_paragraphs)
        if p.style.name == "Heading 2" and p.text == "Luminaires"
    )
    nodes_and_parts_heading_index = next(
        i
        for i, p in enumerate(all_paragraphs)
        if p.style.name == "Heading 2" and p.text == "Nœuds et Parts"
    )
    ascendant_paragraphs = [
        p.text
        for p in all_paragraphs[ascendant_heading_index + 1 : luminaries_heading_index]
        if p.text
    ]
    assert ascendant_paragraphs == expected_ascendant_paragraphs[fixture_name]

    # "Luminaires" recoupé mot pour mot contre le texte réel des deux
    # documents, aux divergences assumées près (voir core.lunation et
    # docgen.builder.add_luminaries_section) : mention systématique de
    # "lumière de secte" (absente du Soleil diurne de Liam dans le document
    # d'origine), énumération systématique des relations d'aspect de la Lune
    # (incomplète dans les deux documents d'origine), combustion du Soleil
    # recalculée depuis `solar_proximity` plutôt que reprise du texte
    # d'Anthony (qui la contredit lui-même : sa propre section "Dignités et
    # réceptions" liste Mercure et Jupiter sous les rayons du Soleil, alors
    # que Luminaires affirme "sous les rayons de personne"), et phase de
    # lunaison classée "Pleine" pour Liam plutôt que "disséminatrice" comme
    # l'affirme son document (bornes vérifiées indépendamment, voir
    # core.lunation).
    expected_luminaries_paragraphs = {
        "anthony": [
            "Soleil : Scorpion, maison 4, régit la maison 1, pérégrin, "
            "sous les rayons de Mercure (13°46' d'écart) et Jupiter (9°09' d'écart), "
            "conjoint à Vénus, Jupiter et la Part de Fortune.",
            "Lune : Lion, maison 1, régit la maison 12, pérégrine, "
            "lumière de secte de ce thème nocturne, conjointe à l'Ascendant et Nœud Sud, "
            "en sextile avec Mars, en carré avec l'amas du Taureau et l'amas du Scorpion, "
            "en trigone avec Mercure, en opposition avec l'amas du Verseau.",
            "Phase de lunaison natale : disséminatrice (écart Soleil-Lune d'environ 270°).",
        ],
        "liam": [
            "Soleil : Balance, maison 9, régit la maison 7, en chute, "
            "lumière de secte de ce thème diurne, sous les rayons de Jupiter (2°18' d'écart), "
            "conjoint à Jupiter et Nœud Sud.",
            "Lune : Taureau, maison 4, régit la maison 6, en exaltation, conjointe à Mars, "
            "en sextile avec l'amas du Cancer, en carré avec l'amas du Lion et l'Ascendant, "
            "en opposition avec Mercure, en aversion avec Nœud Nord, l'amas du Balance et l'amas du Sagittaire.",
            "Phase de lunaison natale : pleine (écart Soleil-Lune d'environ 207°).",
        ],
    }
    luminaries_paragraphs = [
        p.text
        for p in all_paragraphs[luminaries_heading_index + 1 : nodes_and_parts_heading_index]
        if p.text
    ]
    assert luminaries_paragraphs == expected_luminaries_paragraphs[fixture_name]

    # "Nœuds et Parts" recoupé mot pour mot contre le texte réel des deux
    # documents, aux simplifications assumées près (voir
    # docgen.builder.add_nodes_and_parts_section) : mention systématique de
    # la conjonction de signe pour chaque Nœud (`_conjunction_clause`,
    # absente du Nœud Nord d'Anthony dans le document d'origine, qui ne
    # mentionne que le Nœud Sud), proximité réelle au Nœud phrasée sans la
    # nuance "quasiment à l'opposition" du document d'Anthony (le Nœud le
    # plus proche de la Lune est nommé explicitement — Nœud Sud pour
    # Anthony, par conjonction réelle), et configuration d'éclipse reformulée
    # en un gabarit unique donnant les deux écarts chiffrés plutôt que la
    # prose libre (différente entre les deux documents) de "Nœuds et Parts".
    expected_nodes_and_parts_paragraphs = {
        "anthony": [
            "Axe des Nœuds : Nœud Nord en Verseau (maison 7), conjoint à la Part d'Éros, "
            "Nœud Sud en Lion (maison 1), conjoint à l'Ascendant et Lune, "
            "la Lune est proche du Nœud Sud en degrés réels (0°33' d'écart).",
            "La Part de Fortune (Scorpion, maison 4) en opposition avec la Part de l'Esprit (Taureau, maison 10).",
            "Aucune configuration d'éclipse à la naissance : écart Lune-Nœud de 0°33', "
            "écart Soleil-Lune de 89°32' par rapport à la Pleine Lune.",
        ],
        "liam": [
            "Axe des Nœuds : Nœud Nord en Bélier (maison 3), "
            "Nœud Sud en Balance (maison 9), conjoint à Soleil et Jupiter.",
            "La Part de Fortune (Lion, maison 7) et la Part de l'Esprit (Cancer, maison 6) sont en aversion : "
            "aucun aspect ptoléméen ne les relie directement.",
            "Aucune configuration d'éclipse à la naissance : écart Lune-Nœud de 39°11', "
            "écart Soleil-Lune de 26°32' par rapport à la Pleine Lune.",
        ],
    }
    nodes_and_parts_paragraphs = [
        p.text for p in all_paragraphs[nodes_and_parts_heading_index + 1 :] if p.text
    ]
    assert nodes_and_parts_paragraphs == expected_nodes_and_parts_paragraphs[fixture_name]
